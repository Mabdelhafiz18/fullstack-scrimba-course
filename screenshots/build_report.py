from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

SHOTS = r"D:\NU\intern\fullstack-scrimba-course\screenshots"
def shot(name): return os.path.join(SHOTS, name)

doc = Document()

# ---- Base style: Times New Roman 12, single spaced ----
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(12)
pf = normal.paragraph_format
pf.line_spacing = 1.0
pf.space_after = Pt(6)

for s in doc.sections:
    s.top_margin = Inches(1); s.bottom_margin = Inches(1)
    s.left_margin = Inches(1); s.right_margin = Inches(1)

def heading(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(14); r.font.name = "Times New Roman"
    return p

def body(text):
    return doc.add_paragraph(text)

def bullet(text):
    p = doc.add_paragraph(text, style="List Bullet")
    for r in p.runs:
        r.font.name = "Times New Roman"; r.font.size = Pt(12)
    return p

def center_line(text, bold=True, size=12, space=2):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size); r.font.name = "Times New Roman"
    p.paragraph_format.space_after = Pt(space)
    return p

def figure(img, caption, width=5.9):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(shot(img), width=Inches(width))
    c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = c.add_run(caption); rc.italic = True; rc.font.size = Pt(11); rc.font.name = "Times New Roman"
    c.paragraph_format.space_after = Pt(12)

# ================= COVER PAGE =================
logo = os.path.join(SHOTS, "nu-logo.png")
if os.path.exists(logo):
    lp = doc.add_paragraph(); lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lp.add_run().add_picture(logo, width=Inches(1.6))

center_line("Nile University", size=16)
center_line("School of Information Technology and Computer Science", size=13)
center_line("Program of Computer Science", size=13, space=18)

center_line("Student Training Report", size=20, space=4)
center_line("Fullstack Web Development with Scrimba and Coursera", size=13, space=18)

center_line("Submitted in Partial Fulfilment of the Requirements", size=12, space=2)
center_line("For the Bachelor's Degree in Information Technology and Computer Science", size=12, space=2)
center_line("Computer Science", size=12, space=18)

center_line("Submitted by", size=12, space=4)
center_line("Mohamed Abdelhafiz Mohamed  -  221001317", size=12, space=18)

center_line("Training Provider", size=12, space=4)
center_line("Scrimba (Online)  -  Coursera Platform", size=12, space=24)

center_line("Giza - Egypt          Spring 2026", size=12)

doc.add_page_break()

# ================= EXECUTIVE SUMMARY =================
heading("1. Executive Summary")
body("I did an online training course. The name of the course is the Fullstack Development "
     "Specialization. The company Scrimba made this course. The course is on the Coursera website. "
     "The training was about web development.")
body("In this training I learned many new skills. I learned HTML, CSS, and JavaScript. I also "
     "learned React, Node.js, Express, SQL, TypeScript, and Next.js. I used these tools to build "
     "many small projects. I put all my projects in one GitHub repository.")
body("This report explains three things. First, it explains the company. Second, it explains the "
     "training and the projects I built. Third, it explains the skills I learned and the problems I "
     "faced. These skills are close to my study in Computer Science at Nile University. The pictures "
     "of my projects are in the Appendix at the end of the report.")

# ================= THE COMPANY =================
heading("2. The Company")
body("Scrimba is an online school for coding. It started in Norway. Scrimba is not like other "
     "online schools. It uses special videos called 'scrims'. A scrim is a video and code together. "
     "The student can stop the video and change the code inside it. This way of learning is active "
     "and easy to follow.")
body("Coursera is a big website for online learning. Many universities and companies put their "
     "courses on Coursera. Scrimba put its Fullstack Development Specialization on Coursera. So I "
     "could study it there.")
body("The course is big. It has about 22 small courses inside it. It teaches front-end and back-end "
     "web development. The full course needs about three months of study. I studied at my own speed. "
     "At the end, the student builds many projects for a portfolio.")

# ================= THE PROJECT =================
heading("3. The Project")
body("My main work in the training was to build projects. Every project uses one part of the "
     "course. I put all the projects in one GitHub repository. The name of the repository is "
     "'scrimba-fullstack-development-projects'. Below is the list of the projects I built.")

heading("3.1 The Projects I Built")
projects = [
 "Business Card: a simple personal card page. Tools: HTML and CSS.",
 "Space Exploration Page: a landing page about space. Tools: HTML and CSS.",
 "Birthday Website: a fun birthday page with animation. Tools: HTML and CSS.",
 "Basketball Scoreboard: an app to count points for two teams. Tools: JavaScript.",
 "Blackjack Game: a small card game. Tools: JavaScript.",
 "Leads Tracker: a Chrome extension to save links. Tools: JavaScript and localStorage.",
 "API Dashboard: a page that gets live data (time, weather, and a quote). Tools: Fetch and async JavaScript.",
 "Tenzies Game: a dice game. Tools: React.",
 "Assembly Endgame: a word game like Hangman. Tools: React.",
 "Express REST API: a small server for tasks. Tools: Node.js and Express.",
 "SQL Practice: database tables and queries. Tools: SQL.",
 "Next.js Mini Blog: a small blog with many pages. Tools: Next.js and TypeScript.",
 "Full-Stack Blog: a complete blog. The user can add, edit, and delete posts. Tools: Express, SQLite, and React.",
]
for t in projects:
    bullet(t)

heading("3.2 Skills I Gained")
body("I learned how to build web pages that work on a phone and on a computer. I learned how to "
     "write JavaScript to make a page interactive. I learned React to build a page from small "
     "components. I learned how to make a server with Node.js and Express. I learned how to save "
     "data in a SQL database. I learned how to get data from an API on the internet. I also learned "
     "TypeScript and Next.js, which are modern tools for big web apps.")

heading("3.3 Challenges I Faced")
body("Some parts of the training were hard for me. React state was difficult at the start. Async "
     "JavaScript and Promises were also new, so I needed more practice. In the full-stack blog I had "
     "a problem with a port. Port 4000 was already used by another program on my computer. I found "
     "the problem and I fixed it. I moved the backend to port 5050. This problem taught me how to "
     "solve real errors, not only errors in a lesson.")

heading("3.4 Connection to My Study")
body("This training is close to my Computer Science study at Nile University. The programming ideas, "
     "like variables, functions, loops, and conditions, are the same as in my programming courses. "
     "The SQL part is close to my database course. The idea of a client and a server is close to my "
     "web and networks courses. The training gave me hands-on practice for this theory. So the "
     "training and my study help each other.")

# ================= CONCLUSION =================
heading("4. Conclusion")
body("This training was very useful for me. I learned modern web development, from the front-end to "
     "the back-end. I built many real projects. Now I can build a full website with a database.")
body("I liked the active way of learning on Scrimba. The 'scrim' videos helped me a lot. The "
     "hardest parts were React and async JavaScript, but practice made them clear for me.")
body("My advice for other students is simple. Build a small project after every lesson. Do not only "
     "watch the videos. Also save your work on GitHub from the first day.")
body("In the future, I want to learn more about testing, security, and cloud deployment. I also want "
     "to add a user login to my full-stack blog. This training is a strong base for my career as a "
     "software developer.")

# ================= APPENDIX =================
doc.add_page_break()
heading("Appendix: Screenshots of the Projects")
body("This appendix shows pictures of the projects. Projects 10 (Express REST API) and 11 (SQL "
     "Practice) do not have a screen, because they are a server and a database. The other projects "
     "are shown below.")

figs = [
 ("00-portfolio.png", "Figure 1: The portfolio home page. It lists all the projects."),
 ("01-business-card.png", "Figure 2: Project 1 - Business Card (HTML, CSS)."),
 ("02-space.png", "Figure 3: Project 2 - Space Exploration landing page (HTML, CSS)."),
 ("03-birthday.png", "Figure 4: Project 3 - Birthday Website (HTML, CSS)."),
 ("04-scoreboard.png", "Figure 5: Project 4 - Basketball Scoreboard (JavaScript)."),
 ("05-blackjack.png", "Figure 6: Project 5 - Blackjack Game (JavaScript)."),
 ("06-leads-tracker.png", "Figure 7: Project 6 - Leads Tracker (JavaScript, localStorage)."),
 ("07-api-dashboard.png", "Figure 8: Project 7 - API Dashboard with live data (Fetch, async)."),
 ("08-tenzies.png", "Figure 9: Project 8 - Tenzies dice game (React)."),
 ("09-assembly.png", "Figure 10: Project 9 - Assembly Endgame word game (React)."),
 ("12-nextjs-blog-home.png", "Figure 11: Project 12 - Next.js Mini Blog, home page (Next.js, TypeScript)."),
 ("12-nextjs-blog-post.png", "Figure 12: Project 12 - Next.js Mini Blog, a post page."),
 ("14-blog-home.png", "Figure 13: Project 14 - Full-Stack Blog, home page (Express, SQLite, React)."),
 ("14-blog-post.png", "Figure 14: Project 14 - Full-Stack Blog, a post with Markdown."),
 ("14-blog-new.png", "Figure 15: Project 14 - Full-Stack Blog, write a new post."),
]
for img, cap in figs:
    if os.path.exists(shot(img)):
        figure(img, cap)

out = r"D:\NU\intern\fullstack-scrimba-course\Student_Report_Mohamed_Abdelhafiz.docx"
doc.save(out)
print("SAVED:", out)
